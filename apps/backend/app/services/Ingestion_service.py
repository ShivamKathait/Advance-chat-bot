
import asyncio
import fitz
import csv
import uuid
import openpyxl
import pytesseract
import tiktoken

from app.core.exceptions import (
    AppError,
    DocumentParsingError,
    DocumentTooLargeError,
    EmbeddingServiceError,
    UnsupportedFileTypeError,
)
from app.core.logging import logger_adapter
from app.utils.enums import FileType
from app.core.config import settings

from typing import Any, Dict, List
from io import BytesIO, StringIO
from docx import Document
from charset_normalizer import from_bytes
from pdf2image import convert_from_bytes
from google import genai
from tenacity import retry, stop_after_attempt, wait_exponential
from google.genai import types as genai_types


class DocumentParser:
    """Parse different document formats"""

    @staticmethod
    async def parse_pdf(file: bytes) -> str:
        try:
            doc = fitz.open(stream=file, filetype="pdf")

            if len(doc) > settings.MAX_PDF_PAGES:
                raise DocumentTooLargeError(
                    f"PDF has {len(doc)} pages; maximum allowed is {settings.MAX_PDF_PAGES}"
                )

            page_texts: Dict[int, list[str]] = {}
            empty_pages: list[int] = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                parts: list[str] = []

                # Extract tables before get_text so table content isn't duplicated
                for table in page.find_tables().tables:
                    rows = [
                        " | ".join(str(cell) if cell else "" for cell in row)
                        for row in table.extract()
                    ]
                    if rows:
                        parts.append("[Table]\n" + "\n".join(rows))

                text = page.get_text()
                if text.strip():
                    parts.append(text)
                else:
                    empty_pages.append(page_num)

                page_texts[page_num] = parts

            # OCR fallback for scanned / image-only pages
            if empty_pages:
                images = convert_from_bytes(file)
                for page_num in empty_pages:
                    if page_num < len(images):
                        ocr_text = pytesseract.image_to_string(images[page_num])
                        if ocr_text.strip():
                            page_texts[page_num].append(ocr_text)

            result = []
            for page_num in range(len(doc)):
                parts = page_texts[page_num]
                if parts:
                    result.append(f"[Page {page_num + 1}]\n" + "\n".join(parts))

            return "\n\n".join(result)
        except AppError:
            raise
        except Exception as e:
            logger_adapter.error(f"Error parsing PDF: {e}")
            raise DocumentParsingError(f"Failed to parse PDF: {e}")

    @staticmethod
    async def parse_docx(file: bytes) -> str:
        try:
            doc = Document(BytesIO(file))
            text_content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            return "\n".join(text_content)
        except Exception as e:
            logger_adapter.error(f"Error parsing DOCX: {str(e)}")
            raise DocumentParsingError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def decode_text(file_bytes: bytes) -> str:
        result = from_bytes(file_bytes).best()

        if result is None:
            raise DocumentParsingError("Unable to detect file encoding")

        return str(result)

    @staticmethod
    async def parse_txt(file: bytes) -> str:
        try:
            return DocumentParser.decode_text(file)
        except AppError:
            raise
        except Exception as e:
            logger_adapter.error(f"Error parsing Text file: {str(e)}")
            raise DocumentParsingError(f"Failed to parse text file: {e}")

    @staticmethod
    async def parse_markdown(file: bytes) -> str:
        return DocumentParser.decode_text(file)

    @staticmethod
    async def parse_csv(file: bytes) -> str:
        try:
            text = DocumentParser.decode_text(file)
            reader = csv.DictReader(StringIO(text))
            rows = []
            for row in reader:
                rows.append(
                    ", ".join(f"{k}: {v}" for k, v in row.items())
                )
            return "\n".join(rows)
        except AppError:
            raise
        except Exception as e:
            logger_adapter.error(f"Error parsing CSV: {str(e)}")
            raise DocumentParsingError(f"Failed to parse CSV: {e}")

    @staticmethod
    async def parse_xlsx(file: bytes) -> str:
        try:
            wb = openpyxl.load_workbook(BytesIO(file), read_only=True, data_only=True)
            sections = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        rows.append(", ".join(str(c) if c is not None else "" for c in row))
                if rows:
                    sections.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
            return "\n\n".join(sections)
        except Exception as e:
            logger_adapter.error(f"Error parsing XLSX: {str(e)}")
            raise DocumentParsingError(f"Failed to parse XLSX: {e}")

    @staticmethod
    async def parse_document(file: bytes, file_type: str) -> str:
        """
        Parse document based on file type

        Args:
            file: Raw file bytes
            file_type: File extension (.pdf, .docx, .txt, etc.)

        Returns:
            Extracted text
        """
        file_type = file_type.lower()

        if file_type == FileType.PDF:
            return await DocumentParser.parse_pdf(file)
        elif file_type == FileType.DOCX:
            return await DocumentParser.parse_docx(file)
        elif file_type == FileType.TXT:
            return await DocumentParser.parse_txt(file)
        elif file_type in [FileType.MD, FileType.MARKDOWN]:
            return await DocumentParser.parse_markdown(file)
        elif file_type == FileType.CSV:
            return await DocumentParser.parse_csv(file)
        elif file_type == FileType.XLSX:
            return await DocumentParser.parse_xlsx(file)
        else:
            raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")


class TextChunker:
    """Chunk text into manageable pieces"""

    @staticmethod
    def _split_recursive(
        text: str,
        chunk_size: int,
        separators: List[str]
    ) -> List[str]:
        """
        Pure recursive splitting — no overlap applied. Returns a flat list of
        pieces each <= chunk_size, splitting on the best available separator
        (paragraphs -> lines -> spaces -> characters) and recursing into any
        oversized piece with the next separator level.
        """
        separator = separators[-1]
        new_separators: List[str] = []
        for i, sep in enumerate(separators):
            if sep == "":
                separator = sep
                new_separators = []
                break
            if sep in text:
                separator = sep
                new_separators = separators[i + 1:]
                break

        splits = text.split(separator) if separator != "" else list(text)

        # Greedily merge small splits into chunks of <= chunk_size
        final_chunks: List[str] = []
        current_splits: List[str] = []
        current_len = 0

        for split in splits:
            split_len = len(split)
            if split_len > chunk_size:
                if current_splits:
                    final_chunks.append(separator.join(current_splits))
                    current_splits = []
                    current_len = 0
                final_chunks.extend(
                    TextChunker._split_recursive(split, chunk_size, new_separators or [""])
                )
            else:
                sep_len = len(separator) if current_splits else 0
                if current_len + sep_len + split_len > chunk_size and current_splits:
                    final_chunks.append(separator.join(current_splits))
                    current_splits = []
                    current_len = 0
                current_splits.append(split)
                current_len += (len(separator) if len(current_splits) > 1 else 0) + split_len

        if current_splits:
            final_chunks.append(separator.join(current_splits))

        return final_chunks

    @staticmethod
    def chunk_recursive(
        text: str,
        chunk_size: int = None,
        chunk_overlap: int = None,
        separators: List[str] = None
    ) -> List[str]:
        """
        Recursively chunk text using separators

        Tries to split on best separator first:
        1. Double newlines (paragraphs)
        2. Single newlines (lines)
        3. Spaces (words)
        4. Characters

        Args:
            text: Text to chunk
            chunk_size: Max chunk size (default from settings)
            chunk_overlap: Overlap between chunks (default from settings)
            separators: List of separators to try

        Returns:
            List of text chunks
        """

        chunk_size = chunk_size or settings.CHUNK_SIZE
        chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        separators = separators or ["\n\n", "\n", " ", ""]

        # Split first (recursing into oversized pieces), with NO overlap applied
        # at any recursion level — overlap is applied exactly once below, on the
        # final flattened list. Applying it inside the recursion too would double
        # up the prepended tail on chunks produced from an oversized split.
        final_chunks = TextChunker._split_recursive(text, chunk_size, separators)

        if chunk_overlap > 0:
            result = []
            for i, chunk in enumerate(final_chunks):
                if i > 0:
                    tail = final_chunks[i - 1][-chunk_overlap:]
                    chunk = tail + chunk
                result.append(chunk)
        else:
            result = final_chunks

        return [chunk.strip() for chunk in result if chunk.strip()]


    @staticmethod
    def _split_long_chunk(
        text: str,
        chunk_size: int,
        overlap: int
    ) -> List[str]:
        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap

        return chunks

    @staticmethod
    def chunk_by_tokens(
        text: str,
        max_tokens: int = 1000
    ) -> List[str]:
        """
        Chunk text by approximate token count

        Args:
            text: Text to chunk
            max_tokens: Max tokens per chunk

        Returns:
            List of chunks
        """
        try:
            enc = tiktoken.encoding_for_model(settings.OPENAI_MODEL)
            tokens = enc.encode(text)

            chunks = []
            for i in range(0, len(tokens), max_tokens):
                chunk_tokens = tokens[i:i + max_tokens]
                chunk_text = enc.decode(chunk_tokens)
                chunks.append(chunk_text)

            return [c.strip() for c in chunks if c.strip()]

        except Exception as e:
            logger_adapter.error(f"Token-based chunking failed: {str(e)}")
            # Fallback to character-based chunking
            return TextChunker.chunk_recursive(
                text,
                chunk_size=int(max_tokens * 4)  # Rough estimate: 1 token ≈ 4 chars
            )



class EmbeddingGenerator:
    """Generate embeddings for text chunks"""

    def __init__(self):
        self.model = settings.GEMINI_EMBEDDING_MODEL
        self.client = genai.Client(api_key=settings.GEMINI_API_KEY)

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        reraise=True,
    )
    async def _embed_single(self, text: str) -> List[float]:
        response = await self.client.aio.models.embed_content(
            model=self.model,
            contents=text,
        )
        return response.embeddings[0].values

    async def _embed_batch(self, batch: List[str]) -> List[List[float]]:
        return list(await asyncio.gather(*[self._embed_single(t) for t in batch]))

    async def generate_embeddings(
        self,
        texts: List[str],
        batch_size: int = 10
    ) -> List[List[float]]:
        """
        Generate embeddings for multiple texts

        Args:
            texts: List of texts to embed
            batch_size: Batch size for API calls

        Returns:
            List of embedding vectors
        """
        try:
            embeddings = []

            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                batch_embeddings = await self._embed_batch(batch)
                embeddings.extend(batch_embeddings)
                logger_adapter.info(f"Generated embeddings for batch {i // batch_size + 1}")

            return embeddings

        except Exception as e:
            logger_adapter.error(f"Error generating embeddings: {str(e)}")
            raise EmbeddingServiceError(f"Failed to generate embeddings: {e}")

    async def generate_embedding(self, text: str) -> List[float]:
        """
        Generate embedding for single text

        Args:
            text: Text to embed

        Returns:
            Embedding vector
        """
        embeddings = await self.generate_embeddings([text])
        return embeddings[0]


class DocumentIngestionService:
    """
    Main service for document ingestion
    Orchestrates parsing, chunking, and embedding
    """

    def __init__(self):
        self.parser = DocumentParser()
        self.chunker = TextChunker()
        self.embedder = EmbeddingGenerator()
        self._genai_client = genai.Client(api_key=settings.GEMINI_API_KEY)

    async def _generate_chunk_context(
        self, document_text: str, chunk: str, filename: str
    ) -> str:
        """
        Generate a one-sentence anchoring label for a chunk using the full document text.
        This label is prepended before embedding to improve retrieval precision
        (Anthropic Contextual Retrieval, 2024).
        """
        prompt = (
            f"Document: {filename}\n\n"
            f"Full document (first 3000 chars):\n{document_text[:3000]}\n\n"
            f"Chunk:\n{chunk}\n\n"
            "In one sentence, describe which section/topic this chunk belongs to. "
            "Be specific (e.g. 'This is from the Sick Leave section describing entitlement duration.'). "
            "Output only the sentence."
        )
        try:
            response = await self._genai_client.aio.models.generate_content(
                model=settings.GEMINI_MODEL,
                contents=[{"role": "user", "parts": [{"text": prompt}]}],
                config=genai_types.GenerateContentConfig(max_output_tokens=80),
            )
            return response.text.strip()
        except Exception:
            return ""

    async def ingest_document(
            self,
            file: bytes,
            file_type: str,
            document_id: str,
            metadata: Dict[str, Any] = None
            ) -> Dict[str, Any]:
        """
        Ingest a document end-to-end

        Args:
            file: document file bytes
            document_id: Unique document ID
            file_type: File extension (.pdf, .docx, etc.)
            metadata: Additional metadata

        Returns:
            Ingestion result with chunks and embeddings
        """

        metadata = metadata or {}

        logger_adapter.info("Starting document ingestion", document_id=document_id, file_type=file_type)

        try:
            # 1. Parse document
            logger_adapter.info("Parsing document", document_id=document_id)
            text = await self.parser.parse_document(file, file_type)
            logger_adapter.info("Document parsed", document_id=document_id, text_length=len(text))

            # 2. Chunk text
            logger_adapter.info("Chunking text", document_id=document_id)
            chunks = self.chunker.chunk_recursive(text, settings.CHUNK_SIZE, settings.CHUNK_OVERLAP)
            logger_adapter.info("Text chunked", document_id=document_id, num_chunks=len(chunks))

            if len(chunks) > settings.MAX_CHUNKS:
                raise DocumentTooLargeError(
                    f"Document produces {len(chunks)} chunks; maximum allowed is {settings.MAX_CHUNKS}"
                )

            # 3. Contextual enrichment — prepend a topic label to each chunk before embedding
            #    so embeddings carry section context (e.g. "Sick Leave" not just "8 days").
            #    Enabled via CONTEXTUAL_ENRICHMENT_ENABLED config flag; requires re-ingestion.
            filename = metadata.get("filename", "document")
            if settings.CONTEXTUAL_ENRICHMENT_ENABLED:
                logger_adapter.info("Generating chunk context labels", document_id=document_id)
                sem = asyncio.Semaphore(5)  # max 5 concurrent Gemini calls

                async def enrich(chunk: str) -> str:
                    async with sem:
                        label = await self._generate_chunk_context(text, chunk, filename)
                        return f"{label}\n\n{chunk}" if label else chunk

                embed_chunks = list(await asyncio.gather(*[enrich(c) for c in chunks]))
            else:
                embed_chunks = chunks

            # 4. Generate embeddings (over enriched text when enabled)
            logger_adapter.info("Generating embeddings", document_id=document_id, num_chunks=len(chunks))
            embeddings = await self.embedder.generate_embeddings(embed_chunks)
            logger_adapter.info("Embeddings generated", document_id=document_id, num_embeddings=len(embeddings))

            results = {
                "document_id": document_id,
                "file_type": file_type,
                "total_text_length": len(text),
                "num_chunks": len(chunks),
                "chunks": [
                    {
                        "id": str(uuid.uuid5(uuid.UUID(str(document_id)), f"chunk_{i}")),
                        "text": chunk,  # original text shown in sources
                        "embedding": embeddings[i],
                        "metadata": {
                            **metadata,
                            "document_id": str(document_id),
                            "chunk_index": i,
                            "chunk_size": len(chunk),
                            "contextual_enrichment": settings.CONTEXTUAL_ENRICHMENT_ENABLED,
                        }
                    }
                    for i, chunk in enumerate(chunks)
                ]
            }
            logger_adapter.info("Document ingestion completed", document_id=document_id, num_chunks=len(chunks))

            return results
        except Exception as e:
            logger_adapter.error(
                "Document ingestion failed",
                document_id=document_id,
                error=str(e)
            )
            raise

    async def ingest_batch(
        self,
        files: List[tuple],  # List of (file_bytes, document_id, file_type)
        metadata: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """
        Ingest multiple documents

        Args:
            files: List of (file_bytes, document_id, file_type) tuples
            metadata: Shared metadata for all documents

        Returns:
            List of ingestion results
        """
        results = []

        for file_bytes, document_id, file_type in files:
            try:
                result = await self.ingest_document(
                    file_bytes,
                    file_type,
                    document_id,
                    metadata
                )
                results.append(result)
            except Exception as e:
                logger_adapter.error(
                    "Failed to ingest file",
                    document_id=document_id,
                    error=str(e)
                )
                results.append({
                    "document_id": document_id,
                    "status": "failed",
                    "error": str(e)
                })

        return results


# Export service
ingestion_service = DocumentIngestionService()
